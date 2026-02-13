"""Tests for all format parsers."""

from __future__ import annotations

from pathlib import Path

import pytest

from bookworm.parsers.base import get_parser


# ── TXT Parser ─────────────────────────────────────


class TestTxtParser:
    def test_simple_text(self, tmp_path: Path):
        f = tmp_path / "test.txt"
        f.write_text("Hello world.\n\nSecond paragraph.")
        content = get_parser(f).parse(f)
        assert content.metadata.format == "txt"
        assert len(content.chapters) >= 1
        total_paras = sum(len(ch.paragraphs) for ch in content.chapters)
        assert total_paras == 2

    def test_chapter_detection(self, tmp_path: Path):
        f = tmp_path / "chapters.txt"
        f.write_text(
            "Chapter 1 Introduction\n\nContent of chapter one.\n\n"
            "Chapter 2 Main Body\n\nContent of chapter two.\n\n"
            "Chapter 3 Conclusion\n\nFinal thoughts."
        )
        content = get_parser(f).parse(f)
        assert len(content.chapters) >= 3

    def test_chinese_chapter_detection(self, tmp_path: Path):
        f = tmp_path / "cn.txt"
        f.write_text(
            "第一章 开始\n\n第一段内容。\n\n"
            "第二章 中间\n\n第二段内容。\n\n"
            "第三章 结尾\n\n第三段内容。"
        )
        content = get_parser(f).parse(f)
        assert len(content.chapters) >= 3

    def test_empty_file(self, tmp_path: Path):
        f = tmp_path / "empty.txt"
        f.write_text("")
        content = get_parser(f).parse(f)
        assert len(content.chapters) == 0

    def test_no_chapters_splits_into_sections(self, tmp_path: Path):
        f = tmp_path / "long.txt"
        paras = "\n\n".join([f"Paragraph number {i}." for i in range(120)])
        f.write_text(paras)
        content = get_parser(f).parse(f)
        assert len(content.chapters) >= 2  # Should split into chunks


# ── Markdown Parser ────────────────────────────────


class TestMarkdownParser:
    def test_headings_as_chapters(self, tmp_path: Path):
        f = tmp_path / "doc.md"
        f.write_text("# Intro\n\nHello.\n\n# Part Two\n\nWorld.")
        content = get_parser(f).parse(f)
        assert content.metadata.format == "md"
        assert len(content.chapters) == 2
        assert content.chapters[0].title == "Intro"
        assert content.chapters[1].title == "Part Two"

    def test_strip_formatting(self, tmp_path: Path):
        f = tmp_path / "fmt.md"
        f.write_text("# Test\n\n**bold** and *italic* and `code`")
        content = get_parser(f).parse(f)
        para = content.chapters[0].paragraphs[0]
        assert "**" not in para
        assert "*" not in para  # should be stripped
        assert "`" not in para
        assert "bold" in para
        assert "italic" in para

    def test_strip_links(self, tmp_path: Path):
        f = tmp_path / "links.md"
        f.write_text("# Test\n\nVisit [Google](https://google.com) now.")
        content = get_parser(f).parse(f)
        para = content.chapters[0].paragraphs[0]
        assert "Google" in para
        assert "https://" not in para

    def test_strip_images(self, tmp_path: Path):
        f = tmp_path / "img.md"
        f.write_text("# Test\n\n![alt text](image.png)\n\nSome text.")
        content = get_parser(f).parse(f)
        paras = content.chapters[0].paragraphs
        assert not any("image.png" in p for p in paras)

    def test_no_headings(self, tmp_path: Path):
        f = tmp_path / "nohead.md"
        f.write_text("Just a paragraph.\n\nAnother one.")
        content = get_parser(f).parse(f)
        assert len(content.chapters) >= 1

    def test_h2_as_chapters(self, tmp_path: Path):
        f = tmp_path / "h2.md"
        f.write_text("## Section A\n\nText A.\n\n## Section B\n\nText B.")
        content = get_parser(f).parse(f)
        assert len(content.chapters) == 2


# ── DOCX Parser ────────────────────────────────────


class TestDocxParser:
    def test_parse_simple_docx(self, tmp_path: Path):
        """Create a minimal DOCX and parse it."""
        from docx import Document

        doc = Document()
        doc.add_heading("Chapter One", level=1)
        doc.add_paragraph("First paragraph of chapter one.")
        doc.add_paragraph("Second paragraph.")
        doc.add_heading("Chapter Two", level=1)
        doc.add_paragraph("Content of chapter two.")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        content = get_parser(f).parse(f)
        assert content.metadata.format == "docx"
        assert len(content.chapters) == 2
        assert content.chapters[0].title == "Chapter One"
        assert len(content.chapters[0].paragraphs) == 2
        assert content.chapters[1].title == "Chapter Two"
        assert len(content.chapters[1].paragraphs) == 1

    def test_no_headings_single_chapter(self, tmp_path: Path):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Just some text.")
        doc.add_paragraph("More text.")
        f = tmp_path / "nohead.docx"
        doc.save(str(f))

        content = get_parser(f).parse(f)
        assert len(content.chapters) == 1
        assert len(content.chapters[0].paragraphs) == 2


# ── PDF Parser ─────────────────────────────────────


class TestPdfParser:
    def test_parse_simple_pdf(self, tmp_path: Path):
        """Create a minimal PDF and parse it."""
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello World", fontsize=12)
        page.insert_text((72, 100), "Second line of text.", fontsize=12)
        f = tmp_path / "test.pdf"
        doc.save(str(f))
        doc.close()

        content = get_parser(f).parse(f)
        assert content.metadata.format == "pdf"
        assert len(content.chapters) >= 1
        total_paras = sum(len(ch.paragraphs) for ch in content.chapters)
        assert total_paras >= 1


# ── EPUB Parser ────────────────────────────────────


class TestEpubParser:
    def test_parse_minimal_epub(self, tmp_path: Path):
        """Create a minimal EPUB and parse it."""
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier("test123")
        book.set_title("Test Book")
        book.set_language("en")
        book.add_author("Test Author")

        c1 = epub.EpubHtml(title="Chapter 1", file_name="ch1.xhtml", lang="en")
        c1.content = "<html><body><h1>Chapter 1</h1><p>First paragraph.</p><p>Second paragraph.</p></body></html>"
        book.add_item(c1)

        c2 = epub.EpubHtml(title="Chapter 2", file_name="ch2.xhtml", lang="en")
        c2.content = (
            "<html><body><h1>Chapter 2</h1><p>Chapter two content.</p></body></html>"
        )
        book.add_item(c2)

        book.toc = [
            epub.Link("ch1.xhtml", "Chapter 1", "ch1"),
            epub.Link("ch2.xhtml", "Chapter 2", "ch2"),
        ]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ["nav", c1, c2]

        f = tmp_path / "test.epub"
        epub.write_epub(str(f), book)

        content = get_parser(f).parse(f)
        assert content.metadata.title == "Test Book"
        assert content.metadata.author == "Test Author"
        assert content.metadata.format == "epub"
        assert len(content.chapters) >= 2

        # Check paragraphs extracted — nav may be chapter 0, so search all chapters
        all_paras = [p for ch in content.chapters for p in ch.paragraphs]
        assert any("First paragraph" in p for p in all_paras)


# ── get_parser routing ─────────────────────────────


class TestGetParser:
    def test_supported_extensions(self, tmp_path: Path):
        for ext in [".txt", ".md", ".epub", ".pdf", ".docx"]:
            f = tmp_path / f"test{ext}"
            f.touch()
            parser = get_parser(f)
            assert parser is not None

    def test_unsupported_extension(self, tmp_path: Path):
        f = tmp_path / "test.xyz"
        f.touch()
        with pytest.raises(ValueError, match="Unsupported format"):
            get_parser(f)

    def test_mobi_extension(self, tmp_path: Path):
        for ext in [".mobi", ".azw", ".azw3"]:
            f = tmp_path / f"test{ext}"
            f.touch()
            parser = get_parser(f)
            assert parser is not None
