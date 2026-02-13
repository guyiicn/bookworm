"""DOCX parser using python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from bookworm.library.models import Book, BookContent, Chapter

from .base import BaseParser


class DocxParser(BaseParser):
    SUPPORTED_EXTENSIONS = (".docx",)

    def parse(self, file_path: Path) -> BookContent:
        doc = Document(str(file_path))

        # Extract metadata from core properties
        props = doc.core_properties
        title = props.title or file_path.stem
        author = props.author or "Unknown"

        meta = Book(
            id=Book.make_id(str(file_path)),
            file_path=str(file_path),
            title=title,
            author=author,
            format="docx",
            file_size=file_path.stat().st_size,
        )

        # Split into chapters by heading styles
        chapters: list[Chapter] = []
        toc: list[tuple[int, str]] = []
        current_paragraphs: list[str] = []
        current_title = "Document"

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            text = para.text.strip()

            if style_name.startswith("Heading") and text:
                # Save previous chapter
                if current_paragraphs:
                    ch = Chapter(
                        index=len(chapters),
                        title=current_title,
                        paragraphs=current_paragraphs,
                    )
                    chapters.append(ch)
                    toc.append((ch.index, current_title))
                current_title = text
                current_paragraphs = []
            elif text:
                cleaned = re.sub(r"\s+", " ", text)
                current_paragraphs.append(cleaned)

        # Don't forget the last chapter
        if current_paragraphs:
            ch = Chapter(
                index=len(chapters),
                title=current_title,
                paragraphs=current_paragraphs,
            )
            chapters.append(ch)
            toc.append((ch.index, current_title))

        # If no headings found, make one chapter
        if not chapters:
            all_paras = [
                re.sub(r"\s+", " ", p.text.strip())
                for p in doc.paragraphs
                if p.text.strip()
            ]
            if all_paras:
                chapters.append(Chapter(index=0, title=title, paragraphs=all_paras))
                toc.append((0, title))

        meta.total_chapters = len(chapters)
        return BookContent(metadata=meta, chapters=chapters, toc=toc)
